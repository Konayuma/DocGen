import os
import tempfile
from pathlib import Path
from typing import Tuple

try:
    import fitz
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

from docx import Document
from PIL import Image

try:
    import pytesseract
    HAS_TESSERACT = True
except (ImportError, ModuleNotFoundError):
    HAS_TESSERACT = False

from docgen.utils import get_file_extension, safe_read_file


class TextExtractor:
    """Service for extracting text from various document formats."""
    
    @staticmethod
    def extract(filepath: str) -> Tuple[str, dict]:
        """
        Extract text from a document file.
        
        Returns:
            Tuple[str, dict]: (extracted_text, metadata)
            metadata includes: pages, chars, extraction_method, encoding
        """
        ext = get_file_extension(filepath).lower()
        
        if ext == ".pdf":
            return TextExtractor._extract_pdf(filepath)
        elif ext == ".docx":
            return TextExtractor._extract_docx(filepath)
        elif ext == ".txt":
            return TextExtractor._extract_txt(filepath)
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
            return TextExtractor._extract_image(filepath)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    @staticmethod
    def _extract_pdf(filepath: str) -> Tuple[str, dict]:
        """Extract text from PDF using fitz."""
        try:
            if not HAS_FITZ:
                raise ImportError("fitz not available")
            
            doc = fitz.open(filepath)
            text_parts = []
            page_count = len(doc)
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            
            doc.close()
            full_text = "\n".join(text_parts)
            
            # If no text extracted, try OCR
            if not full_text.strip():
                return TextExtractor._extract_pdf_ocr(filepath)
            
            metadata = {
                "pages": page_count,
                "chars": len(full_text),
                "extraction_method": "pdf_native",
            }
            return full_text, metadata
        except Exception:
            # Fallback to OCR
            return TextExtractor._extract_pdf_ocr(filepath)
    
    @staticmethod
    def _extract_docx(filepath: str) -> Tuple[str, dict]:
        """Extract text from DOCX file."""
        try:
            doc = Document(filepath)
            text_parts = []
            
            # Extract paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n".join(text_parts)
            
            metadata = {
                "pages": len(doc.paragraphs) // 50 + 1,  # Rough estimate
                "chars": len(full_text),
                "extraction_method": "docx_native",
            }
            return full_text, metadata
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {str(e)}")
    
    @staticmethod
    def _extract_txt(filepath: str) -> Tuple[str, dict]:
        """Extract text from plain text file."""
        try:
            text = safe_read_file(filepath)
            
            metadata = {
                "pages": len(text) // 3000 + 1,  # Rough estimate: ~3000 chars per page
                "chars": len(text),
                "extraction_method": "txt_native",
            }
            return text, metadata
        except Exception as e:
            raise ValueError(f"TXT extraction failed: {str(e)}")
    
    @staticmethod
    def _extract_image(filepath: str) -> Tuple[str, dict]:
        """Extract text from image using OCR (Tesseract)."""
        if not HAS_TESSERACT:
            return "[Image text extraction requires Tesseract OCR. Install with: choco install tesseract]", {
                "pages": 1,
                "chars": 0,
                "extraction_method": "image_ocr_missing",
                "image_size": None,
            }
        
        try:
            img = Image.open(filepath)
            
            # Preprocess image for better OCR: grayscale, contrast
            img = img.convert("L")
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(img)
            
            metadata = {
                "pages": 1,
                "chars": len(text),
                "extraction_method": "image_ocr",
                "image_size": img.size,
            }
            return text, metadata
        except Exception as e:
            raise ValueError(f"Image extraction failed: {str(e)}")
